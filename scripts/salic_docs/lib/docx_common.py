"""Estilo e helpers compartilhados pelos dois geradores DOCX (dicionário
técnico e catálogo gestor) — identidade visual GovHub (gov-hub.io), extraída
de github.com/GovHub-br/GovHub-skills/tree/main/01-govhub/govhub-visual-identity:
roxo #7A34F3 como cor-assinatura, laranja #F97316 como acento pontual (uso
esparso), fundos claros neutros, fonte Inter.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Paleta oficial GovHub (ver references/palette.md e tokens.css do repo acima)
PRIMARY_PURPLE = RGBColor(0x7A, 0x34, 0xF3)  # cor-assinatura
SECONDARY_PURPLE = RGBColor(0x8B, 0x5C, 0xF6)
PURPLE_600 = RGBColor(0x7C, 0x3A, 0xAD)  # hover/foco
PURPLE_700 = RGBColor(0x5B, 0x21, 0xB6)  # roxo p/ texto pequeno sobre branco (contraste AA)
ACCENT_ORANGE = RGBColor(0xF9, 0x73, 0x16)  # acento pontual — usar com parcimônia
TEXT_STRONG = RGBColor(0x20, 0x20, 0x20)
TEXT_BODY = RGBColor(0x2D, 0x37, 0x48)
TEXT_MUTED = RGBColor(0x66, 0x66, 0x66)
TEXT_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BG_SUBTLE = "F8F9FA"  # zebra de tabela — hex sem # para uso em shading (w:shd)
BORDER_LIGHT = "E2E2E2"

FONT_FAMILY = "Inter"
ASSETS_DIR = Path(__file__).resolve().parents[1] / "assets"
LOGO_PRIMARY = ASSETS_DIR / "logo_horizontal_primary.png"


def _shade_cell(cell: Any, hex_color: str) -> None:
    """Aplica cor de fundo sólida a uma célula de tabela (w:shd, sem
    equivalente direto na API pública do python-docx)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_run_color(cell: Any, rgb: RGBColor, *, bold: bool = False, size: Pt | None = None) -> None:
    for p in cell.paragraphs:
        if not p.runs:
            p.add_run("")
        for r in p.runs:
            r.font.color.rgb = rgb
            r.bold = bold
            if size is not None:
                r.font.size = size


def style_header_row(table: Any, row_index: int = 0) -> None:
    """Cabeçalho de tabela no padrão GovHub: fundo roxo, texto branco em negrito."""
    for cell in table.rows[row_index].cells:
        _shade_cell(cell, "7A34F3")
        _set_run_color(cell, TEXT_WHITE, bold=True)


def zebra_stripe(table: Any, header_rows: int = 1) -> None:
    """Listras alternadas nas linhas de dado (cinza claro `--bg-subtle`)."""
    for i, row in enumerate(table.rows[header_rows:]):
        if i % 2 == 1:
            for cell in row.cells:
                _shade_cell(cell, BG_SUBTLE)


def set_column_widths(table: Any, widths_in: list[float]) -> None:
    """Define larguras fixas por coluna (polegadas). O python-docx exige
    setar em cada célula, não só em `table.columns[i]`, para respeitar no
    Word/LibreOffice."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
    for i, col in enumerate(table.columns):
        if i < len(widths_in):
            col.width = Inches(widths_in[i])


def finalize_table(table: Any, *, header_rows: int = 1, font_size: Pt = Pt(9.5)) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_header_row(table, 0)
    zebra_stripe(table, header_rows)
    for row in table.rows[header_rows:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                if not p.runs:
                    continue
                for r in p.runs:
                    r.font.size = font_size
                    if r.font.color.rgb is None:
                        r.font.color.rgb = TEXT_BODY


def _apply_brand_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_FAMILY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_BODY

    heading_sizes = {0: Pt(30), 1: Pt(20), 2: Pt(15), 3: Pt(12.5)}
    for level, size in heading_sizes.items():
        style_name = "Title" if level == 0 else f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT_FAMILY
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = PRIMARY_PURPLE if level <= 1 else PURPLE_700
        # O template padrão do Word cravava um filete azul de tema
        # (w:pBdr bottom, accent1) no estilo Title — fora da paleta GovHub.
        pPr = style.element.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is not None:
            pPr.remove(pBdr)


def new_document(title: str, subtitle: str, *, landscape: bool = False) -> Document:
    doc = Document()
    _apply_brand_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = section.right_margin = Inches(0.7)

    if LOGO_PRIMARY.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.add_run().add_picture(str(LOGO_PRIMARY), width=Inches(2.6))
        doc.add_paragraph()

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filete laranja (acento pontual GovHub) separando título de subtítulo.
    rule_p = doc.add_paragraph()
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_run = rule_p.add_run("—" * 3)
    rule_run.font.color.rgb = ACCENT_ORANGE
    rule_run.bold = True

    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = TEXT_MUTED
        run.font.size = Pt(12)

    doc.add_page_break()
    return doc


def add_toc(doc: Document) -> None:
    doc.add_heading("Sumário", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    note = doc.add_paragraph(
        "(Clique com o botão direito e escolha \"Atualizar campo\" ao abrir no Word/LibreOffice)"
    )
    note.runs[0].italic = True
    note.runs[0].font.color.rgb = TEXT_MUTED
    doc.add_page_break()


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value if value else "—"
        _set_run_color(row.cells[0], PURPLE_700, bold=True, size=Pt(9.5))
        _set_run_color(row.cells[1], TEXT_BODY, size=Pt(9.5))
        _shade_cell(row.cells[0], BG_SUBTLE)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.color.rgb = TEXT_BODY


def fmt_pct(value: Any) -> str:
    if value is None:
        return "—"
    return f"{value:.2f}%"


def fmt_num(value: Any) -> str:
    if value is None:
        return "—"
    try:
        return f"{int(value):,}".replace(",", ".")
    except (TypeError, ValueError):
        return str(value)


def truncate(text: Any, max_len: int = 400) -> str:
    if text is None:
        return "—"
    text = str(text)
    return text if len(text) <= max_len else text[: max_len - 1] + "…"
