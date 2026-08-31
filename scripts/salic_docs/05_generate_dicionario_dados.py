"""Gera o Dicionário de Dados (DOCX técnico) a partir de output/merged.json —
uma seção por tabela do bronze, agrupadas por domínio (prefixo), com
estrutura semi-templada (não é prosa manual por tabela).

Uso:
    poetry run python scripts/salic_docs/05_generate_dicionario_dados.py [--tables t1,t2] [--out caminho.docx]
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx.shared import Pt

from scripts.salic_docs.lib.docx_common import (
    add_bullet_list,
    add_kv_table,
    add_toc,
    finalize_table,
    fmt_num,
    fmt_pct,
    new_document,
    set_column_widths,
    truncate,
)
from scripts.salic_docs.lib.semantics import DOMINIOS, display_value

OUTPUT_DIR = Path(__file__).resolve().parent / "output"
MERGED_PATH = OUTPUT_DIR / "merged.json"
FINAL_DIR = Path(__file__).resolve().parents[2] / "dbt" / "minc" / "docs" / "salic"
DEFAULT_OUT = FINAL_DIR / "dicionario_dados_salic.docx"

MAX_REL_SHOWN = 12


def _origem_dados(entry: dict[str, Any]) -> str:
    if entry["presente_no_dicionario_original"]:
        return (
            "Sistema SALIC (SQL Server) — schema de origem "
            f"'{entry['prefixo']}', confirmado no dicionário de dados original (SchemaSpy)."
        )
    return (
        "Não identificável com segurança: tabela ausente do dicionário de dados "
        "original do SALIC. Presume-se sistema SALIC pelo schema bronze/prefixo "
        f"'{entry['prefixo']}', mas isso é inferência, não confirmação."
    )


def _add_table_section(doc, entry: dict[str, Any]) -> None:
    doc.add_heading(f"{entry['nome_tabela']}", level=2)
    if entry.get("nome_tabela_origem"):
        p = doc.add_paragraph(f"Nome original (sistema de origem): {entry['nome_tabela_origem']}")
        p.runs[0].italic = True

    doc.add_paragraph(entry.get("descricao") or "Descrição não documentada.")

    add_kv_table(
        doc,
        [
            ("Domínio/assunto", DOMINIOS.get(entry["prefixo"], entry["prefixo"])),
            ("Finalidade", entry.get("descricao") or "Não documentada."),
            ("Quantidade de registros (bronze, hoje)", fmt_num(entry["linhas_atuais_bronze"])),
            (
                "Quantidade de registros (snapshot dicionário original)",
                fmt_num(entry.get("linhas_snapshot_dicionario_original")),
            ),
            ("Quantidade de colunas", str(entry["quantidade_colunas"])),
            ("Origem dos dados", _origem_dados(entry)),
            ("Periodicidade/atualização", "Não documentada na origem nem inferível deste levantamento."),
            (
                "Chave primária (documentada na origem)",
                ", ".join(entry.get("chave_primaria_documentada") or []) or "Não documentada.",
            ),
            (
                "Presente no dicionário de dados original do SALIC",
                "Sim" if entry["presente_no_dicionario_original"] else "Não",
            ),
        ],
    )

    if entry.get("regras_de_negocio_documentadas"):
        doc.add_heading("Regras de negócio documentadas (check constraints)", level=3)
        add_bullet_list(
            doc,
            [f"{c['nome']}: {c['expressao']}" for c in entry["regras_de_negocio_documentadas"]],
        )

    fks = [c for c in entry["colunas"] if c.get("e_chave_estrangeira_documentada")]
    if fks:
        doc.add_heading("Chaves estrangeiras / relacionamentos (documentados na origem)", level=3)
        items = []
        for c in fks[:MAX_REL_SHOWN]:
            for ref in c.get("referencias_documentadas") or []:
                items.append(ref)
        if len(fks) > MAX_REL_SHOWN:
            items.append(f"... e mais {len(fks) - MAX_REL_SHOWN} coluna(s) com FK documentada.")
        add_bullet_list(doc, items or ["Ver colunas individuais abaixo."])

    if entry.get("observacao"):
        doc.add_heading("Observações e limitações", level=3)
        add_bullet_list(doc, entry["observacao"])

    doc.add_heading("Colunas", level=3)
    table = doc.add_table(rows=1, cols=8)
    headers = ["Coluna", "Tipo/Tamanho", "Nulos", "Vazios", "Distintos", "Chave", "Significado", "Exemplo/Domínio observado"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text

    for col in entry["colunas"]:
        row = table.add_row().cells
        row[0].text = col["name"]
        tipo = col.get("tipo_documentado") or "não documentado"
        tam = col.get("tamanho_documentado")
        row[1].text = f"{tipo}" + (f" ({tam})" if tam else "")

        if col.get("ausente_do_perfil_atual"):
            row[2].text = row[3].text = row[4].text = "sem perfil"
        else:
            row[2].text = f"{fmt_num(col.get('null_count'))} ({fmt_pct(col.get('null_pct'))})"
            row[3].text = fmt_num(col.get("empty_count"))
            dist = col.get("distinct_count")
            fonte = col.get("distinct_count_fonte")
            row[4].text = f"{fmt_num(dist)}" + (" (estimado)" if fonte == "estimado_pg_stats" else "")

        flags = []
        if col.get("e_chave_primaria_documentada"):
            flags.append("PK")
        if col.get("e_chave_estrangeira_documentada"):
            flags.append("FK")
        row[5].text = ", ".join(flags) or "—"

        significado = col.get("comentario_original")
        if not significado:
            significado = "[não documentado]"
        row[6].text = truncate(significado, 200)

        if col.get("value_frequency"):
            vals = [
                f"{display_value(col['name'], v['value'])} ({v['freq']}x)"
                for v in col["value_frequency"][:6]
            ]
            row[7].text = truncate("; ".join(vals), 250)
        elif col.get("min_text") is not None or col.get("max_text") is not None:
            row[7].text = truncate(
                f"min: {display_value(col['name'], col.get('min_text'))} / "
                f"max: {display_value(col['name'], col.get('max_text'))}",
                200,
            )
        else:
            row[7].text = "—"

    finalize_table(table, font_size=Pt(8.5))
    set_column_widths(table, [1.0, 1.1, 1.0, 0.7, 1.0, 0.6, 2.2, 2.0])
    doc.add_page_break()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--tables", type=str, default=None)
    parser.add_argument("--out", type=str, default=None)
    args = parser.parse_args()

    merged: dict[str, Any] = json.loads(MERGED_PATH.read_text())
    if args.tables:
        wanted = [t.strip() for t in args.tables.split(",")]
        merged = {k: v for k, v in merged.items() if k in wanted}

    doc = new_document(
        "Dicionário de Dados — SALIC",
        "Schema bronze — camada bruta de dados do Sistema de Apoio às Leis de Incentivo à Cultura\n"
        "Documento técnico — geração automatizada e reprodutível",
        landscape=True,
    )
    add_toc(doc)

    by_prefix: dict[str, list[dict[str, Any]]] = {}
    for entry in merged.values():
        by_prefix.setdefault(entry["prefixo"], []).append(entry)

    for prefix in ["sac", "tabelas", "agentes", "controledeacesso", "bdcorporativo"]:
        entries = by_prefix.get(prefix)
        if not entries:
            continue
        doc.add_heading(f"Domínio: {prefix}", level=1)
        doc.add_paragraph(DOMINIOS.get(prefix, ""))
        doc.add_paragraph(f"{len(entries)} tabela(s) neste domínio.")
        for entry in sorted(entries, key=lambda e: e["nome_tabela"]):
            _add_table_section(doc, entry)

    out_path = Path(args.out) if args.out else DEFAULT_OUT
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"DOCX gerado: {out_path} ({len(merged)} tabelas)")


if __name__ == "__main__":
    main()
